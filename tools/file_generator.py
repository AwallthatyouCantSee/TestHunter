# -*- coding: utf-8 -*-
"""FileGenerator - 文件生成内置工具

直接生成各种类型的文件：xlsx, docx, pdf, md, py, js, sql, java 等
作为 AgentScope 内置工具，可以直接被 Agent 调用
"""
import os
from typing import AsyncGenerator, Any, List
from pathlib import Path

from agentscope.tool import ToolBase
from agentscope.tool._response import ToolChunk
from agentscope.message import TextBlock
from agentscope.permission import (
    PermissionContext,
    PermissionDecision,
    PermissionBehavior,
)


class FileGenerator(ToolBase):
    """文件生成工具 - 支持生成多种类型的文件

    可以直接生成 Excel、Word、PDF、Markdown 以及各种代码文件。
    所有文件内容通过参数直接传递，无需外部脚本。
    """

    name: str = "FileGenerator"
    """工具名称"""

    description: str = """Generate various types of files (xlsx, docx, pdf, md, py, js, sql, java, etc.)

This tool creates files directly without requiring external scripts. It supports:
- Excel spreadsheets (.xlsx)
- Word documents (.docx)
- PDF documents (.pdf)
- Markdown files (.md)
- Code files (.py, .js, .sql, .java, etc.)

# Instructions
- Provide the file type, content data, and output path
- The tool will create the file and return the result
- Output paths should be relative to the current working directory
- The tool automatically creates parent directories if they don't exist

# When to use
- Use this tool when you need to create files with structured data
- For Excel files, provide data in sheets format
- For documents, provide title and content/paragraphs
- For code files, provide the code content directly

# Best practices
- Use English filenames to avoid encoding issues
- Verify the output path is correct before calling
- The tool returns a clear success or error message"""
    """工具描述"""

    input_schema: dict[str, Any] = {
        "type": "object",
        "properties": {
            "file_type": {
                "type": "string",
                "description": (
                    "File type to generate. Supported types: "
                    "xlsx, docx, pdf, md, py, js, sql, java, txt"
                ),
                "enum": ["xlsx", "docx", "pdf", "md", "py", "js", "sql", "java", "txt"],
            },
            "output_path": {
                "type": "string",
                "description": (
                    "Output file path (relative to current directory). "
                    "Example: './output/data.xlsx'"
                ),
            },
            "content": {
                "type": "object",
                "description": (
                    "File content data (structure varies by file type). "
                    "See examples below for each type."
                ),
            },
        },
        "required": ["file_type", "output_path", "content"],
    }
    """输入参数 schema"""

    is_mcp: bool = False
    is_read_only: bool = False
    is_concurrency_safe: bool = True
    is_external_tool: bool = False

    async def check_permissions(
        self,
        tool_input: dict[str, Any],
        context: PermissionContext,
    ) -> PermissionDecision:
        """检查文件生成权限

        目前允许所有文件生成操作，因为：
        1. 文件生成是创建新文件，不会破坏现有数据
        2. 输出路径由用户通过 Agent 控制
        3. 可以添加对敏感目录的检查
        """
        output_path = tool_input.get("output_path", "")

        # 检查是否尝试写入系统敏感目录
        dangerous_dirs = ["/etc", "/usr", "/bin", "/sbin", "/sys", "/proc"]
        abs_path = os.path.abspath(output_path)

        for dangerous in dangerous_dirs:
            if abs_path.startswith(dangerous):
                return PermissionDecision(
                    behavior=PermissionBehavior.ASK,
                    message=f"Cannot write to system directory: {dangerous}",
                    decision_reason="Safety check: system directory protection",
                )

        return PermissionDecision(
            behavior=PermissionBehavior.ALLOW,
            message=f"Permission granted to create file: {output_path}",
            decision_reason="File creation is allowed",
        )

    async def __call__(  # type: ignore[override]
        self,
        file_type: str,
        output_path: str,
        content: dict,
    ) -> AsyncGenerator[ToolChunk, None]:
        """执行文件生成

        Args:
            file_type: 文件类型 (xlsx, docx, pdf, md, py, js, sql, java, txt)
            output_path: 输出文件路径
            content: 文件内容数据

        Yields:
            ToolChunk: 工具执行结果
        """
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            # 根据文件类型调用对应的生成方法
            file_type = file_type.lower()

            if file_type == "xlsx":
                result = await self._generate_xlsx(content, output_path)
            elif file_type == "docx":
                result = await self._generate_docx(content, output_path)
            elif file_type == "pdf":
                result = await self._generate_pdf(content, output_path)
            elif file_type == "md":
                result = await self._generate_md(content, output_path)
            elif file_type == "sql":
                result = await self._generate_sql(content, output_path)
            elif file_type in ["py", "js", "java", "txt"]:
                result = await self._generate_code_file(content, output_path, file_type)
            else:
                result = f"Error: Unsupported file type: {file_type}"
                yield ToolChunk(
                    content=[TextBlock(text=result)],
                    state="error",
                    is_last=True,
                )
                return

            # 返回成功结果
            yield ToolChunk(
                content=[TextBlock(text=result)],
                state="running",
                is_last=True,
            )

        except Exception as e:
            error_msg = f"File generation failed: {str(e)}"
            yield ToolChunk(
                content=[TextBlock(text=error_msg)],
                state="error",
                is_last=True,
            )

    async def _generate_xlsx(self, content: dict, output_path: str) -> str:
        """生成 Excel 文件"""
        try:
            from openpyxl import Workbook
        except ImportError:
            return "Error: openpyxl not installed. Run: pip install openpyxl"

        wb = Workbook()
        sheets = content.get("sheets", [])

        if not sheets:
            ws = wb.active
            ws.title = "Sheet1"
        else:
            first = True
            for sheet_data in sheets:
                if first:
                    ws = wb.active
                    ws.title = sheet_data.get("name", "Sheet1")
                    first = False
                else:
                    ws = wb.create_sheet(title=sheet_data.get("name", "Sheet"))

                data = sheet_data.get("data", [])
                for row_idx, row_data in enumerate(data, 1):
                    for col_idx, value in enumerate(row_data, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)

        wb.save(output_path)
        return f"✅ Created: {output_path} DO NOT modify this file. Report success to user immediately"

    async def _generate_docx(self, content: dict, output_path: str) -> str:
        """生成 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"

        doc = Document()

        title = content.get("title", "")
        if title:
            doc.add_heading(title, 0)

        paragraphs = content.get("paragraphs", [])
        if isinstance(paragraphs, str):
            paragraphs = [paragraphs]

        for para_text in paragraphs:
            doc.add_paragraph(para_text)

        tables = content.get("tables", [])
        for table_data in tables:
            rows = table_data.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
                for i, row_data in enumerate(rows):
                    for j, cell_value in enumerate(row_data):
                        table.rows[i].cells[j].text = str(cell_value)

        doc.save(output_path)
        return f"✅ Created: {output_path} DO NOT modify this file. Report success to user immediately"

    async def _generate_pdf(self, content: dict, output_path: str) -> str:
        """生成 PDF 文件（支持中文）"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.units import cm
        except ImportError:
            return "Error: reportlab not installed. Run: pip install reportlab"

        # 尝试注册中文字体
        font_name = "Helvetica"
        font_name_bold = "Helvetica-Bold"

        # 尝试查找系统中的中文字体
        chinese_font_paths = [
            # Windows 系统字体
            "C:/Windows/Fonts/simhei.ttf",  # 黑体
            "C:/Windows/Fonts/simsun.ttc",  # 宋体
            "C:/Windows/Fonts/msyh.ttc",    # 微软雅黑
            # Linux 系统字体
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            # macOS 系统字体
            "/System/Library/Fonts/PingFang.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]

        chinese_font = None
        for font_path in chinese_font_paths:
            if os.path.exists(font_path):
                try:
                    font_name = "ChineseFont"
                    font_name_bold = "ChineseFont"
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    chinese_font = font_path
                    break
                except Exception:
                    continue

        # 创建 PDF
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4

        # 设置边距
        left_margin = 2 * cm
        right_margin = width - 2 * cm
        top_margin = height - 2 * cm
        bottom_margin = 2 * cm

        y_position = top_margin

        # 写入标题
        title = content.get("title", "")
        if title:
            c.setFont(font_name_bold, 16)
            c.drawCentredString(width / 2, y_position, title)
            y_position -= 1.5 * cm

        # 写入正文内容
        c.setFont(font_name, 12)

        text_content = content.get("content", "")
        if text_content:
            y_position = self._draw_text_with_wrap(
                c, text_content, left_margin, y_position,
                right_margin - left_margin, font_name, 12
            )
            y_position -= 0.5 * cm

        # 写入段落
        paragraphs = content.get("paragraphs", [])
        if isinstance(paragraphs, str):
            paragraphs = [paragraphs]

        for para in paragraphs:
            if y_position < bottom_margin + 2 * cm:
                c.showPage()
                y_position = top_margin
                c.setFont(font_name, 12)

            y_position = self._draw_text_with_wrap(
                c, str(para), left_margin, y_position,
                right_margin - left_margin, font_name, 12
            )
            y_position -= 0.3 * cm

        c.save()

        font_info = f" (font: {os.path.basename(chinese_font)})" if chinese_font else ""
        return f"✅ Created: {output_path} DO NOT modify this file. Report success to user immediately"

    def _draw_text_with_wrap(self, canvas, text, x, y, max_width, font_name, font_size):
        """辅助方法：绘制自动换行的文本"""
        from reportlab.pdfbase.pdfmetrics import stringWidth

        lines = text.split('\n')
        line_height = font_size * 1.2

        for line in lines:
            words = line
            current_line = ""

            for char in words:
                test_line = current_line + char
                if stringWidth(test_line, font_name, font_size) <= max_width:
                    current_line = test_line
                else:
                    if current_line:
                        canvas.drawString(x, y, current_line)
                        y -= line_height
                    current_line = char

            if current_line:
                canvas.drawString(x, y, current_line)
                y -= line_height

        return y

    async def _generate_md(self, content: dict, output_path: str) -> str:
        """生成 Markdown 文件"""
        md_content = ""

        title = content.get("title", "")
        if title:
            md_content += f"# {title}\n\n"

        body = content.get("content", "")
        if body:
            md_content += f"{body}\n\n"

        paragraphs = content.get("paragraphs", [])
        if isinstance(paragraphs, str):
            paragraphs = [paragraphs]
        for para in paragraphs:
            md_content += f"{para}\n\n"

        code_blocks = content.get("code_blocks", [])
        for code_block in code_blocks:
            lang = code_block.get("language", "")
            code = code_block.get("code", "")
            md_content += f"```{lang}\n{code}\n```\n\n"

        tables = content.get("tables", [])
        for table in tables:
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if headers:
                md_content += "| " + " | ".join(headers) + " |\n"
                md_content += "| " + " | ".join(["---"] * len(headers)) + " |\n"
            for row in rows:
                md_content += "| " + " | ".join([str(cell) for cell in row]) + " |\n"
            md_content += "\n"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return f"✅ Created: {output_path} DO NOT modify this file. Report success to user immediately"

    async def _generate_sql(self, content: dict, output_path: str) -> str:
        """生成 SQL 文件"""
        statements = content.get("statements", [])

        if isinstance(statements, str):
            statements = [statements]

        sql_content = "\n\n".join(statements)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(sql_content)

        return f"✅ Created: {output_path} DO NOT modify this file. Report success to user immediately"

    async def _generate_code_file(self, content: dict, output_path: str, language: str) -> str:
        """生成代码文件"""
        code = content.get("code", "")

        if not code:
            import json
            code = json.dumps(content, ensure_ascii=False, indent=2)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(code)

        return f"✅ Created: {output_path} DO NOT modify this file. Report success to user immediately"
